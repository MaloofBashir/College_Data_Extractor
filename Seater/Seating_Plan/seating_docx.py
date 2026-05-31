from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def generate_seating_docx(centre_label, rooms):
    document = Document()

    if document.paragraphs:
        document.paragraphs[0].text = ""

    for index, room in enumerate(rooms):
        if index > 0:
            document.add_section(WD_SECTION_START.NEW_PAGE)

        heading = document.add_heading(f"Centre No: {centre_label}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        room_heading = document.add_paragraph()
        room_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        room_heading.add_run(f"Room No: {room['name']}").bold = True

        table = document.add_table(rows=room["rows"] + 1, cols=room["columns"])
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for column_index, header in enumerate(room["column_headers"]):
            _set_cell_text(table.cell(0, column_index), header or "-", bold=True)

        for row_index, row_values in enumerate(room["grid"], start=1):
            for column_index, value in enumerate(row_values):
                _set_cell_text(table.cell(row_index, column_index), value or "-")

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()
